

from flask import Flask, request, send_file, render_template
from docx import Document
from docx.shared import Pt
from io import BytesIO

app = Flask(__name__)


@app.route('/')
def index():
    return render_template('form.html')


@app.route('/generate-pv', methods=['POST'])
def generate_pv():
    # Récupération des données du formulaire
    titre = request.form['titre']
    marche_num = request.form['marche_num']
    objet = request.form['objet']
    sous_objet = request.form['sous_objet']  # Récupérer le texte du sous-objet
    commission_reception = request.form['commission_reception']

    # Si "Autre" est sélectionné, récupérer la phrase personnalisée
    if commission_reception == "Autre (précisez ci-dessous)...":
        commission_reception = request.form.get('custom_commission_reception', '').strip()

    conformation = request.form.getlist('conformation[]')  # Récupérer toutes les descriptions de Conformation
    recommandations = request.form.getlist('recommandations[]')  # Récupérer les recommandations multiples

    # Récupérer les données des états présents (noms, activités, signatures)
    noms = request.form.getlist('nom[]')
    activites = request.form.getlist('activite[]')
    signatures = request.form.getlist('signature[]')

    # Créer un document Word
    doc = Document()

    # Ajouter le titre "Procès-Verbal" centré
    title_paragraph = doc.add_paragraph()
    title_paragraph.alignment = 1  # Centrer le texte
    title_run = title_paragraph.add_run("Procès-Verbal")
    title_run.bold = True
    title_run.font.size = Pt(18)

    # Ajouter les informations principales
    doc.add_paragraph(f"Marche N° : {marche_num}")
    doc.add_paragraph(f"Objet : {objet}")

    # Ajouter la section "Était Présents"
    doc.add_heading('Était Présents', level=1)

    # Créer le tableau des participants
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'

    # Ajouter les en-têtes du tableau
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Nom'
    hdr_cells[1].text = 'Activité'
    hdr_cells[2].text = 'Signature'

    # Ajouter les données des participants dans le tableau
    for i in range(len(noms)):
        row_cells = table.add_row().cells
        row_cells[0].text = noms[i]
        row_cells[1].text = activites[i]
        row_cells[2].text = signatures[i]

    # Ajouter un espace avant la section "Commission et Réception"
    doc.add_paragraph()  # Ajouter une ligne vide pour l'espacement

    # Ajouter la phrase de Commission et Réception et ajouter les informations de la section Sous-Objet
    doc.add_paragraph(f"{commission_reception} {sous_objet}.")

    # Ajouter la section Conformation
    if conformation:
        doc.add_heading('Conformation', level=1)
        for description in conformation:
            doc.add_paragraph(description)  # Ajouter chaque description de Conformation

    # Ajouter les recommandations
    if recommandations:  # Vérifier si des recommandations ont été soumises
        doc.add_heading("Recommandations", level=1)
        for recommandation in recommandations:
            doc.add_paragraph(recommandation)  # Ajouter chaque recommandation une par une

    # Sauvegarder le document dans un flux de mémoire
    f = BytesIO()
    doc.save(f)
    f.seek(0)

    return send_file(f, as_attachment=True, download_name='Procès-Verbal.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


if __name__ == '__main__':
    app.run(debug=True)
